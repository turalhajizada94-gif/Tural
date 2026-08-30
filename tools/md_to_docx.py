#!/usr/bin/env python3
"""Convert the knowledge-question answer Markdown into a submittable .docx.

Handles only the subset of Markdown used in the answer document: ATX headings,
paragraphs, unordered lists, blockquotes, pipe tables (with <br> line breaks in
cells) and inline bold/italic/code.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(paragraph, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def fill_cell(cell, text):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    segments = re.split(r"<br\s*/?>", text)
    for index, segment in enumerate(segments):
        if index:
            paragraph = cell.add_paragraph()
        add_runs(paragraph, segment.strip())


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def convert(md_path, docx_path):
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    document = Document()
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            index += 1
            continue

        if stripped.startswith("|"):
            block = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index])
                index += 1
            rows = [split_row(r) for r in block if not is_separator(r)]
            width = max(len(r) for r in rows)
            table = document.add_table(rows=0, cols=width)
            table.style = "Table Grid"
            header_is_blank = all(not c for c in rows[0])
            for row_number, row in enumerate(rows):
                cells = table.add_row().cells
                for column, value in enumerate(row):
                    fill_cell(cells[column], value)
                    if row_number == 0 and not header_is_blank:
                        for paragraph in cells[column].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            document.add_paragraph()
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            paragraph = document.add_heading(level=min(level, 4))
            add_runs(paragraph, heading.group(2))
            index += 1
            continue

        if stripped.startswith("> "):
            paragraph = document.add_paragraph(style="Intense Quote")
            add_runs(paragraph, stripped[2:])
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_runs(paragraph, stripped[2:])
            index += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs(paragraph, stripped)
        index += 1

    document.save(docx_path)


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
